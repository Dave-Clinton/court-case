# users/views.py
from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout
from django.core.mail import send_mail
from django.conf import settings
from .models import Profile
import uuid
from django.utils.html import escape
from django.core.exceptions import ValidationError
from django.urls import reverse
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from django.template.loader import render_to_string
from django.shortcuts import get_object_or_404
from django.utils.http import urlsafe_base64_decode
from django.contrib.auth.models import User
from django.http import HttpResponse
from django.utils.encoding import force_str
from django.core.validators import validate_email
from django.contrib.auth import authenticate, login as auth_login
from django.contrib.auth.tokens import default_token_generator
from django.shortcuts import render, redirect, get_object_or_404
from .models import Profile,CaseFile,PassTicket,Payment
from django.contrib.auth.decorators import user_passes_test
from django.http import HttpResponseForbidden
from django.views.decorators.csrf import csrf_protect  
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from docx import Document
from django.http import HttpResponse








def index (request):
    return render (request,'index.html')

@csrf_protect
def register(request):
    if request.method == 'POST':
        first_name = request.POST['first_name']
        last_name = request.POST['last_name']
        email = request.POST['email']
        password = request.POST['password']
        password2 = request.POST['password2']
        phone_number = request.POST['phone_number']
        gender = request.POST['gender']
        county = request.POST['county']

        if password == password2:
            if User.objects.filter(email=email).exists():
                messages.error(request, 'Email is already taken')
            elif len(password) < 6 or not any(char.isdigit() for char in password) or not any(char.isalpha() for char in password):
                messages.error(request, 'Password must be at least 6 characters long and contain both numbers and letters')
        
            else:
                is_admin = email in settings.ADMIN_EMAILS
                user = User.objects.create_user(username=email, email=email, password=password)  # Create a regular user
                user.save()
                verification_code = uuid.uuid4()
                profile = Profile(
                    user=user, 
                    first_name=first_name, 
                    last_name=last_name, 
                    verification_code=verification_code,
                    phone_number=phone_number, 
                    gender=gender, 
                    county=county
                )
                profile.is_admin = is_admin  # Assign admin status based on email
                profile.save()
                send_verification_email(user, verification_code)
                messages.success(request, 'Registration successful. Please check your email to verify your account.')
                return redirect('verify')
        else:
            messages.error(request, 'Passwords do not match')

    return render(request, 'register.html')



def send_verification_email(user, verification_code):
    uid = urlsafe_base64_encode(force_bytes(user.pk))
    token = str(verification_code)
    verification_link = reverse('verify_email', kwargs={'uidb64': uid, 'token': token})
    verification_url = f"{settings.SITE_URL}{verification_link}"

    subject = 'Verify your email address'
    message = f'Hi {user.username},\n\nPlease click the following link to verify your email address: {verification_url}\n\nThank you!'
    from_email = settings.DEFAULT_FROM_EMAIL
    recipient_list = [user.email]
    send_mail(subject, message, from_email, recipient_list)



def verify_email(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = get_object_or_404(User, pk=uid)
        profile = Profile.objects.get(user=user)
        verification_code = uuid.UUID(token)
        
        if profile.verification_code == verification_code and not profile.is_verified:
            profile.is_verified = True
            profile.save()
            messages.success(request, 'Email verified successfully!')
            return redirect('login')
        else:
            messages.error(request, 'Verification link is invalid or has already been used.')
            return render(request, 'verify.html')
    except (TypeError, ValueError, OverflowError, User.DoesNotExist, Profile.DoesNotExist):
        messages.error(request, 'Verification link is invalid.')
        return render(request, 'verify.html')


def verify (request):
    return render (request, 'verify.html')


@csrf_protect
def login_view(request):
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '').strip()

        # Validate email
        try:
            validate_email(email)
        except ValidationError:
            messages.error(request, 'Invalid email format')
            return render(request, 'login.html')

        # Check if email and password are provided
        if not email or not password:
            messages.error(request, 'Email and password are required')
            return render(request, 'login.html')

        # Authenticate user
        user = authenticate(request, username=email, password=password)
        if user is not None:
            auth_login(request, user)
            return redirect('homeland')  # Adjust the redirect as needed
        else:
            messages.error(request, 'Invalid credentials')

    return render(request, 'login.html')



@csrf_protect
def password_reset_request(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        user = User.objects.filter(email=email).first()
        if user:
            uid = urlsafe_base64_encode(force_bytes(user.pk))
            token = default_token_generator.make_token(user)
            password_reset_link = reverse('password_reset_confirm', kwargs={'uidb64': uid, 'token': token})
            password_reset_url = f"{settings.SITE_URL}{password_reset_link}"

            subject = 'Reset your password'
            message = f'Hi {user.username},\n\nPlease click the following link to reset your password: {password_reset_url}\n\nThank you!'
            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [user.email]
            send_mail(subject, message, from_email, recipient_list)

            messages.success(request, 'A link to reset your password has been sent to your email.')
            return redirect('login')
        else:
            messages.error(request, 'This email is not registered.')
    return render(request, 'password_reset_form.html')


@csrf_protect
def password_reset_confirm(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = get_object_or_404(User, pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user is not None and default_token_generator.check_token(user, token):
        if request.method == 'POST':
            password = request.POST.get('password')
            password2 = request.POST.get('password2')
            if password == password2:
                if len(password) < 8 or not any(char.isdigit() for char in password) or not any(char.isalpha() for char in password) or not any(char.islower() for char in password) or not any(char.isupper() for char in password):
                    messages.error(request, 'Password must be at least 8 characters long, contain both numbers and letters, and have both lowercase and uppercase letters')
                else:
                    user.set_password(password)
                    user.save()
                    messages.success(request, 'Password has been reset successfully!')
                    return redirect('login')
            else:
                messages.error(request, 'Passwords do not match')
        return render(request, 'password_reset_confirm.html')
    else:
        messages.error(request, 'The reset link is invalid, possibly because it has already been used.')
        return redirect('password_reset_form')
    

def logout_view(request):
    logout(request)
    return redirect('login')


#edit profile
@login_required
def edit_profile(request):
    profile = request.user.profile

    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        phone_number = request.POST.get('phone_number')
        gender = request.POST.get('gender')
        county = request.POST.get('county')

        # Update the profile fields
        profile.first_name = first_name
        profile.last_name = last_name
        profile.phone_number = phone_number
        profile.gender = gender
        profile.county = county
        profile.save()

        messages.success(request, 'Your profile was successfully updated!')
        return redirect('homeland')  # Redirect to the home page after updating

    return render(request, 'edit_profile.html', {'profile': profile})

#delete profile
@login_required
def delete_profile(request):
    user = request.user

    if request.method == 'POST':
        user.delete()
        messages.success(request, 'Your account has been deleted.')
        return redirect('register')  

    return render(request, 'delete_profile.html')


#home data delivery view

#def homeland(request):
    case_files = CaseFile.objects.filter(user=request.user)
    
    
    completed_cases = CaseFile.objects.filter(user=request.user, status='completed')
    case_tickets = []
    for case in completed_cases:
        tickets = PassTicket.objects.filter(case_file=case)
        payments = Payment.objects.filter(ticket__in=tickets)
        case_tickets.append({
            'case': case,
            'tickets': tickets,
            'payments': payments,
        })
    
    context = {
        'case_files': case_files,
        'case_tickets': case_tickets,
    }
    
    return render(request, 'home.html', context)





@login_required
def homeland(request):
    # Retrieve query parameters
    case_id = request.GET.get('case_id', '')
    first_name = request.GET.get('first_name', '')
    form_type = request.GET.get('form_type', '')
    user_type = request.GET.get('user_type', '')

    # Filter case files based on query parameters
    case_files = CaseFile.objects.filter(user=request.user)
    if case_id:
        case_files = case_files.filter(id__icontains=case_id)
    if first_name:
        case_files = case_files.filter(user__profile__first_name__icontains=first_name)
    if form_type:
        case_files = case_files.filter(station__icontains=form_type)
    if user_type:
        case_files = case_files.filter(user_type__icontains=user_type)

    # Filter completed cases and gather associated tickets and payments
    completed_cases = case_files.filter(status='completed')
    case_tickets = []
    for case in completed_cases:
        tickets = PassTicket.objects.filter(case_file=case)
        payments = Payment.objects.filter(ticket__in=tickets)
        case_tickets.append({
            'case': case,
            'tickets': tickets,
            'payments': payments,
        })

    context = {
        'case_files': case_files,
        'case_tickets': case_tickets,
    }
    
    return render(request, 'home.html', context)



#ticket details
@login_required
def ticket_detail(request, ticket_id):
    ticket = get_object_or_404(PassTicket, id=ticket_id)
    context = {
        'ticket': ticket,
    }
    return render(request, 'ticket_detail.html', context)

#case filde details
@login_required
def case_detail(request, case_id):
    case = get_object_or_404(CaseFile, id=case_id)
    context = {
        'case': case,
    }
    return render(request, 'case_detail.html', context)


#download ticket and case files as word docx

def export_case_to_word(request, case_id):
    case = get_object_or_404(CaseFile, id=case_id)
    document = Document()
    document.add_heading(f'Case Details for {case.user_type}', 0)

    document.add_paragraph(f'User_type: {case.user_type}')
    document.add_paragraph(f'Postal Address: {case.postal_address}')
    document.add_paragraph(f'Telephone Number: {case.telephone_number}')
    document.add_paragraph(f'Agent: {case.agent}')
    document.add_paragraph(f'Caretaker: {case.caretaker}')
    document.add_paragraph(f'Auctioneer: {case.auctioneer}')
    document.add_paragraph(f'Duration of Stay: {case.duration_of_stay}')
    document.add_paragraph(f'Monthly Rent: {case.monthly_rent}')
    document.add_paragraph(f'Year of Entry: {case.year_of_entry}')
    document.add_paragraph(f'Deposit Paid: {case.deposit_paid}')
    document.add_paragraph(f'Cause of Action: {case.cause_of_action}')
    document.add_paragraph(f'Problem: {case.problem}')
    document.add_paragraph(f'OCS Police Station: {case.ocs_police_station}')
    document.add_paragraph(f'Status: {case.get_status_display()}')
    document.add_paragraph(f'Created At: {case.created_at.strftime("%Y-%m-%d %H:%M:%S")}')
    # Add other relevant details

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=case_{case.id}.docx'
    document.save(response)
    return response


from django.http import FileResponse

def download_ticket_document(request, ticket_id):
    ticket = get_object_or_404(PassTicket, id=ticket_id)
    
    # Check if the document exists
    if ticket.document:
        response = FileResponse(ticket.document.open('rb'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={ticket.document.name}'
        return response
    else:
        # Handle case where the document is not available
        return HttpResponse('No document found for this ticket.', status=404)





#def dashboard(request):
      # Retrieve all case files and corresponding pass tickets
    case_files = CaseFile.objects.all()  # No filtering by user
    pass_tickets = PassTicket.objects.filter(case_file__in=case_files)
    
    context = {
        'case_files': case_files,
        'pass_tickets': pass_tickets,
    }


    return render(request, 'dashboard.html', context)






def dashboard(request):
    # Initialize the queryset
    case_files = CaseFile.objects.all()

    # Get filter parameters from the request
    case_id = request.GET.get('case_id')
    first_name = request.GET.get('first_name', '')
    form_type = request.GET.get('form_type')
    user_type = request.GET.get('user_type')

    # Apply filters if provided
    if case_id:
        case_files = case_files.filter(id=case_id)
    if first_name:
        case_files = case_files.filter(user__profile__first_name__icontains=first_name)
    if form_type:
        case_files = case_files.filter(ocs_police_station__icontains=form_type)
    if user_type:
        case_files = case_files.filter(user_type__icontains=user_type)

    # Retrieve the pass tickets for the filtered case files
    pass_tickets = PassTicket.objects.filter(case_file__in=case_files)
    
    context = {
        'case_files': case_files,
        'pass_tickets': pass_tickets,
    }

    return render(request, 'dashboard.html', context)







@login_required
def details(request):
    if request.method == 'POST':
        input_option = request.POST.get('input_option')
        form_type = request.POST.get('form_type')

        if input_option == 'upload':
            # Handle document upload
            file_upload = request.FILES.get('file_upload')
            if file_upload:
                case_file = CaseFile(
                    user=request.user,
                    file_upload=file_upload,
                    form_type=form_type,
                    status='in_progress',
                    created_at=timezone.now()
                )
                case_file.save()
                return redirect('homeland')  # Adjust the redirect as needed
        else:
            # Handle manual form submission
            user_type = request.POST.get('user_type')
            postal_address = request.POST.get('postal_address')
            telephone_number = request.POST.get('telephone_number')
            agent = request.POST.get('agent', '')
            caretaker = request.POST.get('caretaker', '')
            auctioneer = request.POST.get('auctioneer', '')
            duration_of_stay = request.POST.get('duration_of_stay', '')
            monthly_rent = request.POST.get('monthly_rent', '')
            year_of_entry = request.POST.get('year_of_entry', '')
            deposit_paid = request.POST.get('deposit_paid', '')
            cause_of_action = request.POST.get('cause_of_action', '')
            problem = request.POST.get('problem', '')
            ocs_police_station = request.POST.get('ocs_police_station', '')

            case_file = CaseFile(
                user=request.user,
                user_type=user_type,
                postal_address=postal_address,
                telephone_number=telephone_number,
                agent=agent,
                caretaker=caretaker,
                auctioneer=auctioneer,
                duration_of_stay=duration_of_stay,
                monthly_rent=monthly_rent,
                year_of_entry=year_of_entry,
                deposit_paid=deposit_paid,
                cause_of_action=cause_of_action,
                problem=problem,
                form_type=form_type,
                ocs_police_station=ocs_police_station,
                status='in_progress',
                created_at=timezone.now()

            )
            case_file.save()
            return redirect('homeland')  # Adjust the redirect as needed

    return render(request, 'details.html')





#def create_pass_ticket(request, case_file_id):
    case_file = get_object_or_404(CaseFile, id=case_file_id)
    
    if request.method == 'POST':
        content = request.POST.get('content')
        PassTicket.objects.create(case_file=case_file, content=content)
        case_file.status = 'in_review'
        case_file.save()
        return redirect('dashboard')

    context = {'case_file': case_file}
    return render(request, 'affidavit.html', context)



    return render(request, 'affidavit.html', context)




def create_pass_ticket(request, case_file_id):
    case_file = get_object_or_404(CaseFile, id=case_file_id)
    
    if request.method == 'POST':
        if 'document' in request.FILES:
            document = request.FILES['document']
            PassTicket.objects.create(case_file=case_file, document=document)
            case_file.status = 'in_review'
            case_file.save()
            return redirect('dashboard')
        else:
            return render(request, 'affidavit.html', {
                'case_file': case_file,
                'error': 'Please upload a valid document.',
            })

    context = {'case_file': case_file}
    return render(request, 'affidavit.html', context)



def update_pass_ticket_status(request, ticket_id):
    ticket = get_object_or_404(PassTicket, id=ticket_id)
    action = request.POST.get('action')

    if action == 'complete':
        ticket.case_file.status = 'completed'
    elif action == 'revision':
        ticket.case_file.status = 'in_revision'

    ticket.case_file.save()
    return redirect('review_dashboard')


#def review_dashboard(request):
      # Retrieve all case files and corresponding pass tickets
    case_files = CaseFile.objects.all()  # No filtering by user
    pass_tickets = PassTicket.objects.filter(case_file__in=case_files)
    
    context = {
        'case_files': case_files,
        'pass_tickets': pass_tickets,
    }
    if request.method == 'POST':
        ticket_id = request.POST.get('ticket_id')
        action = request.POST.get('action')
        ticket = get_object_or_404(PassTicket, id=ticket_id)
        case_file = ticket.case_file

        if action == 'complete':
            case_file.status = 'completed'
        elif action == 'revision':
            case_file.status = 'in_revision'

        case_file.save()

    return render(request, 'review_dashboard.html', context)



def review_dashboard(request):
    # Retrieve query parameters
    case_id = request.GET.get('case_id', '')
    first_name = request.GET.get('first_name', '')
    form_type = request.GET.get('form_type', '')
    user_type = request.GET.get('user_type', '')

    # Filter case files based on query parameters
    case_files = CaseFile.objects.all()
    if case_id:
        case_files = case_files.filter(case_id__icontains=case_id)
    if first_name:
        case_files = case_files.filter(user__profile__first_name__icontains=first_name)

    if form_type:
        case_files = case_files.filter(station__icontains=form_type)
    if user_type:
        case_files = case_files.filter(user_type__icontains=user_type)

    # Retrieve pass tickets related to filtered case files
    pass_tickets = PassTicket.objects.filter(case_file__in=case_files)
    
    context = {
        'case_files': case_files,
        'pass_tickets': pass_tickets,
    }

    if request.method == 'POST':
        ticket_id = request.POST.get('ticket_id')
        action = request.POST.get('action')
        ticket = get_object_or_404(PassTicket, id=ticket_id)
        case_file = ticket.case_file

        if action == 'complete':
            case_file.status = 'completed'
        elif action == 'revision':
            case_file.status = 'in_revision'

        case_file.save()

    return render(request, 'review_dashboard.html', context)


#delete ticket
def delete_pass_ticket(request, ticket_id):
    ticket = get_object_or_404(PassTicket, id=ticket_id)
    case_file = ticket.case_file
    
    if request.method == 'POST':
        ticket.delete()
        # Update the status of the associated CaseFile
        case_file.status = 'in_progress'
        case_file.save()
        return redirect('dashboard')
    
    context = {'ticket': ticket}
    return render(request, 'delete_pass_ticket.html', context)






#def edit_pass_ticket(request, ticket_id):
    ticket = get_object_or_404(PassTicket, id=ticket_id)
    if request.method == 'POST':
        content = request.POST.get('content')
        ticket.content = content
        ticket.save()
        return redirect('dashboard')

    context = {'ticket': ticket}
    return render(request, 'edit_pass_ticket.html', context)



@login_required
def edit_pass_ticket(request, ticket_id):
    ticket = PassTicket.objects.get(id=ticket_id)
    if request.method == 'POST':
        # Process the form submission
        new_content = request.POST.get('content')
        ticket.content = new_content
        ticket.save()

        # Redirect to review_dashboard
        return redirect('dashboard')

    context = {'ticket': ticket}
    return render(request, 'edit_pass_ticket.html', context)



#paypal
from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
import paypalrestsdk
import json

paypalrestsdk.configure({
    "mode": settings.PAYPAL_MODE,
    "client_id": settings.PAYPAL_CLIENT_ID,
    "client_secret": settings.PAYPAL_CLIENT_SECRET
})




def make_payment(request, ticket_id):
    # Retrieve the ticket
    ticket = get_object_or_404(PassTicket, id=ticket_id)
    
    # Provide the ticket and any other necessary context to the frontend
    context = {
        'ticket_id': ticket.id,
        'case_file_id': ticket.case_file.id,  # Assuming `case_file` is a related field
    }
    
    return render(request, 'make_payment.html', context)

#def capture_payment(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        order_id = data.get('orderID')

        # Capture the payment
        payment = paypalrestsdk.Payment.find(order_id)
        if payment.execute({"payer_id": data.get('payerID')}):
            # Mark the payment as made in the session
            request.session['payment_made'] = True
            return JsonResponse({'status': 'success', 'message': 'Payment successfully captured'})
        else:
            return JsonResponse({'status': 'error', 'message': 'Payment capture failed'}, status=400)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)





@csrf_exempt
def capture_payment(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        order_id = data.get('orderID')
        payer_id = data.get('payerID')
        ticket_id = data.get('ticketID')  # Assuming you send ticket ID with payment data

        # Retrieve the ticket
        try:
            ticket = PassTicket.objects.get(id=ticket_id)
        except PassTicket.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Ticket not found'}, status=404)

        # Capture the payment
        payment = paypalrestsdk.Payment.find(order_id)
        if payment.execute({"payer_id": payer_id}):
            # Create a Payment record in the database
            Payment.objects.create(
                user=request.user,
                ticket=ticket,  # Use the ticket object
                amount=5.00,  # The amount of the transaction
            )
            # Mark the payment as made in the session
            request.session['payment_made'] = True
            return JsonResponse({'status': 'success', 'message': 'Payment successfully captured'})
        else:
            return JsonResponse({'status': 'error', 'message': 'Payment capture failed'}, status=400)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)

